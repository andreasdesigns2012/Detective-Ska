#!/usr/bin/env python3
"""
Detective Ska — Art Guides (per room) generator
=============================================

Creates a separate DOCX focused ONLY on ART PRODUCTION for every room:
 - what the room is (kitchen/lobby/boiler/darkroom…)
 - where clues are staged (what object holds the clue)
 - what numbers/letters must be visible for puzzles
 - background layers + foreground layers
 - prop lists (static vs interactive vs breakable)
 - lighting palette suggestions (hex) + mood notes
 - decals/grime/storytelling set dressing
 - animation loops + VFX needs
 - Unity 2D technical notes (sorting layers, parallax, collider boundaries)

Requirements:
  pip install python-docx

Usage:
  python generate_detective_ska_art_guides.py
  python generate_detective_ska_art_guides.py --out Detective_Ska_Art_Guides.docx
"""

from __future__ import annotations

import argparse
import datetime as _dt
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Reuse the canonical room list from the GDD generator (single source of truth)
from generate_detective_ska_gdd import build_content, RoomSpec, BuildingSpec  # type: ignore


def add_title_page(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.italic = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Generated on {_dt.date.today().isoformat()}")
    doc.add_page_break()


def add_toc(doc: Document, title: str = "Table of Contents") -> None:
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click here and choose “Update Field” to generate the TOC."
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    doc.add_page_break()


def bullets(doc: Document, items: List[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc: Document, items: List[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def kv_table(doc: Document, rows: List[Tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v


def _palette_for_building(building_code: str) -> Dict[str, str]:
    # Consistent per-building palette so the whole game reads as one world.
    return {
        "B-01": "Base: #2A2F2D (mildew dark) | Neon spill: #3BC7FF (cold cyan) | Accent rust: #7A3D2B | Paper: #D8D2C6",
        "B-02": "Base: #1F2428 (wet asphalt) | Sodium: #FFB000 (dirty yellow) | Red safety: #B0001A | Steel: #8A8F98",
        "B-03": "Base: #141012 (almost black) | Antique wood: #5B3A2E | Candle: #FFCC7A | Dust light: #E9E1D1",
        "HQ":   "Base: #E7E9EC (fluorescent white) | Shadow: #8B939E | Warning: #FF3B30 | Navy: #1C2F4A",
    }


def _room_tags(room: RoomSpec) -> List[str]:
    name = room.name.lower()
    tags: List[str] = []
    if "lobby" in name or "entrance" in name:
        tags.append("lobby")
    if "kitchen" in room.summary.lower() or "kitchen" in name:
        tags.append("kitchen")
    if "boiler" in name or "steam" in room.summary.lower():
        tags.append("boiler")
    if "rooftop" in name or "roof" in room.floor.lower():
        tags.append("rooftop")
    if "dark" in name or "darkroom" in room.summary.lower() or "photo lab" in name.lower():
        tags.append("darkroom")
    if "archive" in name or "records" in name:
        tags.append("archive")
    if "hallway" in name or "corridor" in room.summary.lower():
        tags.append("hallway")
    if "chapel" in name:
        tags.append("chapel")
    if "basement" in room.floor.lower():
        tags.append("basement")
    if "exterior" in room.floor.lower() or "courtyard" in name:
        tags.append("exterior")
    return tags


def _art_blueprint(room: RoomSpec, building_code: str) -> Dict[str, List[str]]:
    """
    A *more detailed* art plan per room. This expands on the GDD art notes with:
    - Layers
    - Prop lists
    - Clue staging specifics (WHAT object must clearly show WHAT text/numbers)
    - Technical Unity notes
    """
    tags = _room_tags(room)

    # Defaults (applies to all rooms)
    layers = [
        "BG_FAR: skyline / distant walls / big silhouettes (low detail, big shapes).",
        "BG_MID: main room walls, large furniture silhouettes, windows, doors.",
        "BG_NEAR: foreground trims, pipes close to camera, hanging objects.",
        "FG: occluders (door frames, pillars, tall props) to create depth.",
        "FX: rain/steam/dust/fog overlays, light cones, sparks.",
    ]

    # Technical baseline so artists + programmers align
    unity_notes = [
        "Use Unity 2D URP lights for key mood lights (1–3 lights max per room).",
        "Sorting Layers (recommended): BG_FAR < BG_MID < BG_NEAR < ACTORS < FG < VFX < UI.",
        "Parallax (if used): BG_FAR 0.15, BG_MID 0.35, BG_NEAR 0.6 (relative camera movement).",
        "Collision: keep major wall colliders simple rectangles; place detail sprites without colliders.",
        "Interactables must have 2 visual states: IDLE and HIGHLIGHT (outline/glow).",
        "Clue props must support a close-up sprite or UI panel texture (clean, readable).",
    ]

    # Room-specific sets
    scene_zones: List[str] = []
    hero_props: List[str] = []
    interactable_visuals: List[str] = []
    clue_staging: List[str] = []
    decals_story: List[str] = []
    anim_loops: List[str] = []
    vfx_needs: List[str] = []

    # Generic expansions by tag
    if "lobby" in tags:
        scene_zones += [
            "Zone A: Entrance door silhouette + rain reflection strip.",
            "Zone B: Mail wall (readable labels) + broken intercom focus.",
            "Zone C: Stairwell door framed as a ‘next step’ composition cue."
        ]
        hero_props += [
            "Mailboxes set (10–20), with 3–5 readable names, the rest scratched.",
            "Intercom panel (broken) with exposed wires (3 wire colors visible).",
            "Window bars casting shadow lines on the floor puddle.",
        ]
        interactable_visuals += [
            "Intercom wiring puzzle: 3 wire ends must be visually distinct (red/white/black).",
            "Mailbox: 1 highlighted mailbox should have a slightly bent door to hint interaction.",
        ]
        clue_staging += [
            "Clue holder: a torn envelope in trash OR behind vending machine; include a readable stamp/time.",
            "If the room contains a matchbook: ‘SAINT VESTA’ text must be readable on close-up."
        ]
        decals_story += [
            "Water stain trails under mailboxes (suggests long-term leaks).",
            "Footprint smears near stairwell (someone came through recently)."
        ]
        anim_loops += ["Neon flicker reflection in puddle (2–3 frame loop)."]
        vfx_needs += ["Rain drip particles near ceiling crack; subtle dust motes in window light cone."]

    if "kitchen" in tags:
        scene_zones += [
            "Kitchen Nook: small counter, sink, upper cabinets, fridge (storytelling clutter).",
            "Living Corner: couch + photo wall (clue focal point)."
        ]
        hero_props += [
            "Kitchen counter set: plates, mug, bottle, ashtray, knife block (no gore).",
            "Fridge with 2–3 magnets (one magnet can hint the next location).",
        ]
        decals_story += [
            "Grease stain around stove area; cigarette burn on table; water ring stains.",
            "One cabinet door slightly open to lead the eye toward an interactable."
        ]

    if "boiler" in tags:
        scene_zones += [
            "Boiler Core: big cylinder silhouette with animated heat shimmer.",
            "Valve Wall: three valves clearly labeled V1/V2/V3 for puzzle readability.",
            "Sealed Door: heavy steel door with pressure gauge nearby."
        ]
        hero_props += [
            "Pipe maze kit (straight, elbow, T-joint) with color bands to guide navigation.",
            "Pressure gauge with readable marks: 70 / 40 / 30 (chalk or engraved).",
            "Control panel with empty fuse slot (very clear)."
        ]
        interactable_visuals += [
            "Valves must rotate (simple 8-frame loop).",
            "Fuse slot highlight state: subtle glow outline when player has fuse.",
        ]
        clue_staging += [
            "Numbers for puzzle must be readable in both normal view and close-up: 70 / 40 / 30.",
            "Steam burst sources must be placed so player learns the hazard pattern."
        ]
        vfx_needs += ["Thick steam particles (layered) + occasional pressure burst (screen-space fog)."]

    if "darkroom" in tags:
        scene_zones += [
            "Red Light Cone: one strong red practical light that defines silhouettes.",
            "Hanging Photo Line: photos are the ‘reward wall’ composition.",
            "Safe Corner: small safe with clear dial/number plate."
        ]
        hero_props += [
            "Photo line set: 6–10 hanging photos, at least 3 with visible wedding-cake numbers (7, 1, 9).",
            "Chemical trays (3–5) with reflective surface.",
            "Red safety lamp (flicker optional)."
        ]
        clue_staging += [
            "Puzzle numbers must be readable: wedding topper numbers ‘7’, ‘1’, ‘9’.",
            "Photo back date stamps (for close-up UI): make 3 distinct stamps (OLD / MID / NEW).",
            "Audio Tape Fragment label must be readable on close-up: “Final Husband”."
        ]
        vfx_needs += [
            "Low visibility red light grading; optional chemical fog overlay if trays break.",
            "Drip animation from hanging photos."
        ]

    if "archive" in tags:
        scene_zones += [
            "Desk Focus: lamp cone, typewriter, case list paper.",
            "Cabinet Wall: file cabinets with consistent labeling system.",
            "Calendar Wall: big readable circled dates."
        ]
        hero_props += [
            "Wall calendar with 3 circled dates: 10, 14, 22 + month abbreviations OCT / APR / FEB.",
            "Desk safe with 3-digit window (needs to show code entry).",
            "Pinned city map with 3–5 colored pins."
        ]
        clue_staging += [
            "Calendar cipher must be readable: 'OCT', 'APR', 'FEB' and circled day numbers.",
            "Case number list paper must be readable in close-up (even if blurred in main view)."
        ]
        anim_loops += ["Desk lamp subtle flicker + dust motes in light cone."]

    if "rooftop" in tags:
        scene_zones += [
            "Skyline Depth: 3 parallax layers (far skyline, mid neon signs, near roof props).",
            "Antenna Focus: composition points player toward antenna.",
            "Edge Safety: low wall, avoid distracting fall risk visuals."
        ]
        hero_props += [
            "Water tank (big silhouette) with hatch; graffiti text 'GULL ST' readable.",
            "Antenna mast with 4-direction rotation markers (subtle)."
        ]
        clue_staging += [
            "Graffiti must clearly show: “GULL ST” and an arrow/compass mark.",
            "Radio static UI overlay optional; but antenna should have a visual 'signal strength' change (sparks/LED)."
        ]
        vfx_needs += ["Rain streak particles + lightning flash (rare).", "Wind-driven fog wisps crossing camera."]

    if "chapel" in tags:
        scene_zones += [
            "Altar Table: rings bowl as center prop.",
            "Chair Rows: cover geometry for combat.",
            "Window Light: rain streak parallax + candle flicker."
        ]
        hero_props += [
            "Ring bowl with exactly 12 readable ring silhouettes (countable).",
            "Candle stand weapon prop (breakable sparks)."
        ]
        clue_staging += [
            "Cabinet lock code clue: note must be readable: “All my husbands in one bowl.”",
            "If player counts rings, ring highlights should pulse when hovered."
        ]

    if "basement" in tags:
        scene_zones += [
            "Hanging Light Cone: main focus with strong falloff.",
            "Body Tableau: Elena silhouette in shadow (tasteful framing).",
            "Drain/Wall Stains: subtle movement (drip)."
        ]
        hero_props += [
            "Hanging lamp (swing slightly).",
            "Final photograph asset (close-up required).",
            "Revolver prop near hand."
        ]
        clue_staging += [
            "Photograph must be crisp in close-up: Ska + Elena smiling.",
        ]
        vfx_needs += ["Dust in light cone; slow camera zoom vignette; subtle drip particles."]

    if "exterior" in tags:
        scene_zones += [
            "Courtyard Drain Center: puddle reflection and composition anchor.",
            "Dumpster Corner: hides torn letter (clue).",
            "Upper Balcony: camera flash source point."
        ]
        hero_props += [
            "Laundry lines with cloth sprites (sway loop).",
            "Camera flash cone (light volume).",
            "Fire escape ladder (climbable)."
        ]
        clue_staging += [
            "Torn letter close-up must show 1–2 readable phrases: “photo lab” and “wedding proofs”."
        ]

    # Always include a "clue prop readability" rule and camera notes
    camera_notes = [
        "Camera composition: 70% of interactables should sit on rule-of-thirds points.",
        "Clue pickups: add a 0.3s pause + close-up overlay so players can read numbers/words.",
        "Avoid over-noisy backgrounds behind clue text (keep text on clean, high-contrast surfaces).",
    ]

    # Add explicit details for a few key rooms where the user asked for “kitchen + clues + nums”
    overrides: Dict[str, Dict[str, List[str]]] = {
        "B1-R3": {
            "scene_zones": [
                "Kitchen Nook (right side): counter + sink + upper cabinets (dirty dishes, mug with lipstick stain).",
                "Photo Wall (left side): 6 framed photos, 1 is the key evidence (scratched face) framed brighter.",
                "Bedroom Door (back): half-open, creates depth and unease."
            ],
            "clue_staging": [
                "Photo Wall: scratches must be strong and readable from mid-distance (the erased husband).",
                "Jewelry box: symbols must be readable: Moon, Key, Eye, Rose, Knife (simple icon set).",
                "Poem clue prop: a small paper under a frame with the line: “Moon watches. Key opens. Eye judges. Rose lies. Knife ends.”",
            ],
            "interactable_visuals": [
                "Jewelry box: 5 ring slots with etched icons (icons must read at 1x scale).",
                "Kitchen drawer: razor blades visible but not gory (silvery highlights).",
            ],
        },
        "B1-R5": {
            "clue_staging": [
                "Calendar must clearly show: OCT / APR / FEB, with circled days 10 / 14 / 22.",
                "Safe keypad close-up: ensure digits are readable (use high contrast).",
                "Case number list sheet: printed, slightly yellow paper; typed mono font."
            ]
        },
        "B2-R4": {
            "clue_staging": [
                "Hanging photos: at least 3 wedding-cake topper numbers visible in-frame: 7, 1, 9.",
                "Microcassette fragment label close-up: “Final Husband – If found, play on full deck.”",
                "Safe dial: numbers clear under red light (use bright paint/engraving)."
            ]
        },
        "B3-R3": {
            "clue_staging": [
                "Ring bowl: rings must be countable (12). Consider a subtle hover highlight per ring.",
                "Note prop: readable line: “All my husbands in one bowl.” (handwritten)."
            ]
        },
    }

    if room.code in overrides:
        o = overrides[room.code]
        scene_zones = o.get("scene_zones", scene_zones)
        hero_props = o.get("hero_props", hero_props)
        interactable_visuals = o.get("interactable_visuals", interactable_visuals)
        clue_staging = o.get("clue_staging", clue_staging)
        decals_story = o.get("decals_story", decals_story)
        anim_loops = o.get("anim_loops", anim_loops)
        vfx_needs = o.get("vfx_needs", vfx_needs)

    # Ensure each section is not empty
    if not scene_zones:
        scene_zones = ["Primary zone: follow the GDD layout; keep 3 clear composition anchors (entry, clue, exit)."]
    if not hero_props:
        hero_props = ["Use the GDD interactables list as required props; add 5–10 ambient props for storytelling."]
    if not clue_staging:
        clue_staging = ["Clue prop must be readable in close-up; use a clean surface and high-contrast text/numbering."]
    if not anim_loops:
        anim_loops = ["Ambient loop: subtle light flicker or rain drip to keep the scene alive."]
    if not vfx_needs:
        vfx_needs = ["Minimum: dust motes/light cone OR subtle fog layer depending on mood."]

    return {
        "Room Tags": [", ".join(tags) if tags else "none"],
        "Building Palette": [_palette_for_building(building_code).get(building_code, "Use game-wide noir palette.")],
        "Scene Zones (What areas exist in the room)": scene_zones,
        "Layer Plan (how to build the background)": layers,
        "Hero Props (must model/draw)": hero_props,
        "Interactable Visual States": interactable_visuals if interactable_visuals else [
            "Each interactable needs: idle sprite + highlight sprite + 'used' sprite if applicable."
        ],
        "Clue Staging (WHAT object shows WHAT text/numbers)": clue_staging,
        "Decals & Storytelling Set Dressing": decals_story if decals_story else [
            "Add 6–12 decals: water stains, scratches, posters, cigarette burns, footprint smears."
        ],
        "Animation Loops (simple, cheap)": anim_loops,
        "VFX Needs": vfx_needs,
        "Camera & Readability Notes": camera_notes,
        "Unity 2D Technical Notes": unity_notes,
        "Naming/Export Conventions": [
            "Background sprites: BG_{roomcode}_{layer}_{desc}.png",
            "Props: PROP_{roomcode}_{desc}.png",
            "Decals: DECAL_{roomcode}_{desc}.png",
            "Puzzle close-ups: UI_CLUE_{roomcode}_{desc}.png",
        ],
    }


def render_room_art_guide(doc: Document, building_code: str, room: RoomSpec) -> None:
    doc.add_heading(f"{room.code} — {room.name} ({room.floor})", level=2)
    doc.add_paragraph(room.summary)

    # Quick “What is this room?” table
    kv_table(doc, [
        ("Room Objective (Design)", room.objective),
        ("Key Puzzle (Design)", " / ".join(room.puzzle[:2]) if room.puzzle else "N/A"),
        ("Primary Clue Rewards (Design)", "; ".join(room.clue_rewards[:2]) if room.clue_rewards else "N/A"),
    ])

    # Artist guide blueprint
    blueprint = _art_blueprint(room, building_code)
    for section, items in blueprint.items():
        doc.add_heading(section, level=3)
        bullets(doc, items)

    doc.add_page_break()


def render(doc: Document) -> None:
    content = build_content()
    buildings: List[BuildingSpec] = content["buildings"]  # type: ignore
    police_hq_rooms: List[RoomSpec] = content["police_hq_rooms"]  # type: ignore

    add_title_page(
        doc,
        "Detective Ska",
        "ART GUIDES — Per Room (Background + Props + Clue Readability + Unity Notes)",
    )
    add_toc(doc)

    doc.add_heading("How to use this document", level=1)
    bullets(doc, [
        "Each room has a production checklist for background artists + prop artists + VFX.",
        "Anything listed under “Clue Staging” MUST be readable in a close-up UI sprite.",
        "Keep silhouettes strong and shapes readable; noir works best when the scene reads in 2 seconds.",
        "Use building palettes so players subconsciously feel progression by color and light quality.",
    ])
    doc.add_page_break()

    # Buildings
    for b in buildings:
        doc.add_heading(f"{b.code} — {b.name}", level=1)
        doc.add_paragraph(b.address_style)
        doc.add_paragraph(b.tone)
        doc.add_heading("Building Color/Lighting Palette", level=2)
        bullets(doc, [_palette_for_building(b.code).get(b.code, "Use noir palette.")])
        doc.add_page_break()

        for room in b.rooms:
            render_room_art_guide(doc, b.code, room)

    # Police HQ
    doc.add_heading("HQ — Police Headquarters (Intermission)", level=1)
    bullets(doc, [_palette_for_building("HQ")["HQ"]])
    doc.add_page_break()
    for room in police_hq_rooms:
        render_room_art_guide(doc, "HQ", room)

    doc.add_heading("Global Asset Checklist (All Rooms)", level=1)
    bullets(doc, [
        "Neon sign kit (3–5 variants), flicker animation.",
        "Rain overlay particles (light, medium, heavy).",
        "Puddle tile set + reflection mask.",
        "Paper props: receipts, letters, photos, diary pages, certificates (UI close-ups).",
        "Cassette deck hero prop + microcassette fragment close-up.",
        "Generic grime decals pack (50+).",
    ])


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Generate Detective Ska Art Guides docx.")
    p.add_argument("--out", default="Detective_Ska_Art_Guides_By_Room.docx", help="Output DOCX file name/path.")
    return p.parse_args()


def main() -> None:
    args = parse_args()
    doc = Document()
    render(doc)
    doc.save(args.out)
    print(f"Saved: {args.out}")
    print("Tip: Open the DOCX in Word and update fields to populate the Table of Contents.")


if __name__ == "__main__":
    main()

