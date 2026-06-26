#!/usr/bin/env python3
"""
Detective Ska — Room Index (quick reference) generator
-----------------------------------------------------

Creates a DOCX that lists EVERY room with:
 - Room code + name
 - Room type label (e.g., Bedroom / Living Room / Lobby / Boiler Room)
 - Bullet points for the most important items (clues/props like newspapers, photos, letters, tapes, etc.)

Requirements:
  pip install python-docx

Usage:
  python generate_detective_ska_room_index.py
  python generate_detective_ska_room_index.py --out Detective_Ska_Room_Index.docx
"""

from __future__ import annotations

import argparse
import datetime as _dt
import re
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Pull the canonical rooms from the GDD generator so everything stays consistent.
from generate_detective_ska_gdd import build_content, RoomSpec, BuildingSpec  # type: ignore


ROOM_TYPES: Dict[str, str] = {
    # Building 1
    "B1-R1": "Lobby",
    "B1-R2": "Closet",
    "B1-R3": "Apartment",
    "B1-R4": "Hallway",
    "B1-R5": "Office",
    "B1-R6": "Rooftop",
    # Building 2
    "B2-R1": "Courtyard",
    "B2-R2": "Boiler Room",
    "B2-R3": "Apartment",
    "B2-R4": "Darkroom",
    # Building 3
    "B3-R1": "Entrance Hall",
    "B3-R2": "Bedroom",
    "B3-R3": "Chapel",
    "B3-R4": "Hallway (Top Floor)",
    "B3-R5": "Dark Room",
    "B3-R6": "Basement",
    # Police HQ
    "HQ-R1": "Lobby",
    "HQ-R2": "Records Room",
}


KEYWORDS = [
    "newspaper",
    "photo",
    "photograph",
    "diary",
    "letter",
    "certificate",
    "ledger",
    "cassette",
    "tape",
    "map",
    "calendar",
    "rings",
    "ring",
    "matchbook",
    "intercom",
    "safe",
    "typewriter",
]


def title_page(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(subtitle).italic = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Generated on {_dt.date.today().isoformat()}")
    doc.add_page_break()


def bullets(doc: Document, items: List[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def extract_codes_and_numbers(room: RoomSpec) -> List[str]:
    # Gather all puzzle text into one string and extract digits.
    blob = " ".join(room.puzzle or [])
    nums = re.findall(r"\b\d+\b", blob)
    # De-duplicate while preserving order.
    seen = set()
    out: List[str] = []
    for n in nums:
        if n not in seen:
            seen.add(n)
            out.append(n)
    return out[:8]


def important_props(room: RoomSpec) -> List[str]:
    combined = " ".join((room.summary or "") + " " + " ".join(room.interactables or []) + " " + " ".join(room.clue_rewards or []))
    combined_l = combined.lower()

    props: List[str] = []
    for kw in KEYWORDS:
        if kw in combined_l:
            # Normalize to a nicer label
            label = kw.capitalize()
            if kw in ("photo", "photograph"):
                label = "Photos / photographs"
            if kw == "newspaper":
                label = "Newspaper articles"
            if kw == "cassette":
                label = "Cassette recorder"
            if kw == "tape":
                label = "Audio tape"
            props.append(label)

    # De-duplicate
    out: List[str] = []
    for p in props:
        if p not in out:
            out.append(p)
    return out[:8]


def room_key_points(room: RoomSpec) -> List[str]:
    pts: List[str] = []

    # Always include objective in a short form
    if room.objective:
        pts.append(f"Objective: {room.objective}")

    # Clues/rewards (top 2)
    if room.clue_rewards:
        pts.append(f"Key clue/reward: {room.clue_rewards[0]}")
        if len(room.clue_rewards) > 1:
            pts.append(f"Additional: {room.clue_rewards[1]}")

    # Important props like photos/newspapers/letters
    props = important_props(room)
    if props:
        pts.append("Important props: " + ", ".join(props))

    # Key numbers/codes
    nums = extract_codes_and_numbers(room)
    if nums:
        pts.append("Key numbers/codes shown in scene: " + ", ".join(nums))

    # Notable interactables (filtered)
    interesting = []
    for it in (room.interactables or []):
        it_l = it.lower()
        if any(k in it_l for k in KEYWORDS):
            interesting.append(it)
    if interesting:
        pts.append("Must-be-clear interactables: " + "; ".join(interesting[:4]))

    # Ensure we always have at least a few bullets
    if len(pts) < 4 and room.encounters:
        pts.append(f"Encounter beat: {room.encounters[0]}")

    return pts[:8]


def render_room(doc: Document, room: RoomSpec) -> None:
    doc.add_heading(f"{room.code} — {room.name}", level=2)
    rtype = ROOM_TYPES.get(room.code, "Room")
    doc.add_paragraph(f"Type: {rtype}")
    bullets(doc, room_key_points(room))


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default="Detective_Ska_Room_Index.docx", help="Output DOCX path/name")
    args = ap.parse_args()

    content = build_content()
    buildings: List[BuildingSpec] = content["buildings"]  # type: ignore
    police_hq_rooms: List[RoomSpec] = content["police_hq_rooms"]  # type: ignore

    doc = Document()
    title_page(
        doc,
        "Detective Ska",
        "Room Index — quick room list with key props/clues",
    )

    doc.add_heading("Rooms by Location", level=1)
    doc.add_paragraph(
        "This is a quick reference for the team. Each room lists the room type and the important items "
        "(clue props like photos, newspapers, letters, certificates, tapes, and any key numbers/codes)."
    )

    for b in buildings:
        doc.add_heading(f"{b.code} — {b.name}", level=1)
        for room in b.rooms:
            render_room(doc, room)

    doc.add_heading("HQ — Police Headquarters", level=1)
    for room in police_hq_rooms:
        render_room(doc, room)

    doc.save(args.out)
    print(f"Saved: {args.out}")


if __name__ == "__main__":
    main()
