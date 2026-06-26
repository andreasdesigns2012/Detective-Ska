#!/usr/bin/env python3
"""
Detective Ska - Game Design Document (GDD) generator
===================================================

This script generates a long, *usable* DOCX document (roughly 15–30+ pages,
depending on your Word settings) for a 2D Unity detective game.

It includes:
 - Story + tone + characters
 - Evidence & clue system
 - Mission flow
 - Three main investigation buildings with room-by-room specs
 - Optional bar-fight sequence (in parentheses as requested)
 - Puzzles, enemies/encounters, dialogue snippets
 - Art tasks (background artist, character artist, VFX, sound)
 - Final cassette twist ending and basement reveal

Requirements
------------
pip install python-docx

Usage
-----
python generate_detective_ska_gdd.py
python generate_detective_ska_gdd.py --out "Detective_Ska_GDD.docx"
python generate_detective_ska_gdd.py --out "docs/Detective_Ska_GDD.docx"
"""

from __future__ import annotations

import argparse
import datetime as _dt
from dataclasses import dataclass, field
from typing import List, Optional, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# -----------------------------
# DOCX helpers
# -----------------------------

def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic


def add_title_page(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, title, bold=True)
    p.runs[0].font.size = 280000  # 28pt in EMU-ish; python-docx uses half-points internally but accepts ints

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p2, subtitle, italic=True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p3, f"Generated on {_dt.date.today().isoformat()}")

    doc.add_page_break()


def add_toc(doc: Document, title: str = "Table of Contents") -> None:
    """
    Insert a TOC field. Word will populate it when the user updates fields:
      Word: References -> Update Table
    """
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()

    # Create: w:fldSimple w:instr="TOC \o "1-3" \h \z \u"
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click here and choose “Update Field” to generate the TOC."
    r.append(t)
    fld.append(r)
    p._p.append(fld)

    doc.add_page_break()


def add_kv_table(doc: Document, rows: List[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = k
        row_cells[1].text = v


def add_bullets(doc: Document, items: List[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc: Document, items: List[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_dialogue_block(doc: Document, lines: List[str]) -> None:
    for line in lines:
        doc.add_paragraph(line, style="Intense Quote")


def add_section_break(doc: Document) -> None:
    doc.add_page_break()


# -----------------------------
# Content model
# -----------------------------

@dataclass
class EvidenceItem:
    code: str
    name: str
    found_in: str
    description: str
    unlocks: str


@dataclass
class RoomSpec:
    code: str
    name: str
    floor: str
    summary: str
    objective: str
    layout: List[str]
    interactables: List[str]
    puzzle: List[str]
    clue_rewards: List[str]
    encounters: List[str]
    fail_states: List[str]
    dialogue: List[str]
    art_background: List[str]
    art_characters: List[str]
    vfx_sfx: List[str]
    exit_condition: str


@dataclass
class BuildingSpec:
    code: str
    name: str
    address_style: str
    tone: str
    rooms: List[RoomSpec] = field(default_factory=list)


@dataclass
class MissionSpec:
    code: str
    name: str
    purpose: str
    overview: str
    player_goals: List[str]
    key_beats: List[str]
    deliverables: List[str]
    notes: List[str] = field(default_factory=list)


# -----------------------------
# GDD content (hand-authored)
# -----------------------------

def build_content() -> Dict[str, object]:
    # Main cast / tone
    story_overview = (
        "Detective Ska is a noir detective story set in a rain-choked city of neon and soot. "
        "Detective Ska (Adrian Ska) is a once-respected investigator whose life collapses after "
        "his wife, Elena, vanishes without a trace. The police stall, the city forgets, and Ska "
        "does what he always does: he hunts the truth alone.\n\n"
        "A rumor overheard in a bar turns grief into a case: a woman matching Elena’s description "
        "has been linked to decades of disappearances across multiple cities. The deeper Ska digs, "
        "the more the evidence suggests an impossible idea—he is not hunting a kidnapper. He is "
        "unknowingly investigating the woman he loved."
    )

    characters = [
        ("Detective Adrian Ska", "Player character. Visually: tired eyes, rumpled coat, half-lit cigarette. "
                                 "Personality: sharp, controlled, self-blaming. Strength: observation. Weakness: obsession."),
        ("Elena (Alias: “Silent Reaper”)", "Ska’s missing wife. Appears in photos across decades. Charm weaponized. "
                                           "She leaves a final cassette confession. The twist: she was the killer all along."),
        ("Frank Ramirez", "Bartender and ex-cop. Knows the city’s secrets. Gives Ska the first lead. "
                          "A moral mirror: he warns Ska what obsession does to men."),
        ("Mara Voss", "A private archivist who sells “truth” to the highest bidder. Not a villain, but dangerous."),
        ("The Grays", "A small crew of hired thugs used to recover evidence. They serve as recurring encounters."),
    ]

    # Gameplay pillars
    pillars = [
        "Investigation Loop: Explore → Observe → Collect Evidence → Solve Puzzle → Unlock Next Space/Lead.",
        "Noir Tone: Rain, neon, cigarettes, cheap whiskey, long shadows, quiet dread.",
        "Evidence Board: Evidence items connect into a “Truth Web” that gates progression.",
        "Room-Based Progression: Each room delivers one concrete clue + one gameplay beat (puzzle/encounter).",
        "Truth Escalation: Early clues look like kidnapping; mid-game suggests serial pattern; late-game confirms Elena’s authorship.",
    ]

    evidence = [
        EvidenceItem("E-01", "Bar Receipt (Black Raven)", "Mission M-01 / Black Raven Bar",
                     "A receipt with a time stamp and a handwritten note: “VESTA knows. Don’t trust the badge.”",
                     "Unlocks Building B-01: Marrow Apartments (lobby door)."),
        EvidenceItem("E-02", "Burned Matchbook", "B-01 / R1 Lobby",
                     "Matchbook from the Saint Vesta Residence with a lipstick mark.",
                     "Unlocks the theory node: “Elena visited Saint Vesta.”"),
        EvidenceItem("E-03", "Scratched Photograph", "B-01 / R3 Apartment 1A",
                     "Photo of Elena with a man; his face is scratched out with violent strokes.",
                     "Unlocks Evidence Board connection: “Pattern of erased husbands.”"),
        EvidenceItem("E-04", "Maintenance Keyring", "B-01 / R2 Janitor Closet",
                     "A ring labeled: Archive, Roof, Boiler, Records.",
                     "Unlocks B-01 upper floors + later gates B-02 Boiler access."),
        EvidenceItem("E-05", "Police Case Number List", "B-01 / R5 Archive Office",
                     "A list of case numbers tied to missing husbands, copied from police archives.",
                     "Unlocks Mission M-02: Police Records Visit."),
        EvidenceItem("E-06", "Wedding Certificate Copy", "Police HQ / Records Room",
                     "Copy of a wedding certificate: Elena’s signature under a different surname.",
                     "Unlocks the “Many Names” node + enables B-02 Rooftop Photo Lab door."),
        EvidenceItem("E-07", "Audio Tape Fragment", "B-02 / R4 Hidden Photo Lab",
                     "A microcassette fragment labeled: “Final Husband – If found, play on full deck.”",
                     "Unlocks B-03 top-floor Dark Room entry condition."),
        EvidenceItem("E-08", "Caretaker Ledger", "B-03 / R2 Caretaker’s Room",
                     "Ledger listing payments from “E.” for private room access and “disposal service.”",
                     "Unlocks basement hatch mechanism."),
        EvidenceItem("E-09", "Cassette Confession", "B-03 / R5 Dark Room (Top Floor)",
                     "Elena’s full confession to the murders, followed by a gunshot.",
                     "Triggers Final Sequence: Basement reveal + ending headline."),
    ]

    missions = [
        MissionSpec(
            "M-00", "Prologue: Rain Without End",
            "Introduce the mood, Ska’s loss, and the obsession theme.",
            "A short playable cutscene establishes Ska’s apartment, the empty side of the bed, and the unanswered phone.",
            ["Walk through Ska’s apartment", "Inspect Elena’s items", "Leave into the rain"],
            [
                "Player sees Elena’s ring box left open (missing ring).",
                "A voicemail plays: static, then a breath, then silence.",
                "Ska grabs his coat, gun, and notebook: the case begins."
            ],
            ["Teach interaction: inspect, collect, note-taking", "Set noir tone: rain + neon + isolation"],
        ),
        MissionSpec(
            "M-01", "Mission 1: The Black Raven Bar",
            "Kick off investigation with the bartender’s intel and the first evidence item.",
            "Ska drinks to forget. Frank Ramirez gives him a rumor he shouldn’t repeat. Ska chooses the case anyway.",
            ["Get Frank to talk", "Eavesdrop on a suspicious stranger", "Leave with a concrete location lead"],
            [
                "Ska arrives. Music low, neon sign flickers.",
                "The Stranger mentions an old building: Marrow Apartments.",
                "Frank warns Ska about Elena’s pattern.",
                "(Optional: the player enters a bar fight sequence after pushing too hard.)",
                "Ska finds a receipt with a hidden note (E-01)."
            ],
            ["Evidence E-01", "Tutorial for Evidence Board + objective tracking"],
            notes=["Keep bar fight optional; if included, put it in parentheses in the doc text as requested."],
        ),
        MissionSpec(
            "M-02", "Intermission: Police HQ Records Visit",
            "Use the police archives to turn scattered clues into a timeline.",
            "Ska breaks protocol and revisits his old workplace to pull sealed records on the missing-husband cases.",
            ["Gain access to Records Room", "Copy a wedding certificate", "Avoid being seen by Internal Affairs"],
            [
                "Stealth-lite: avoid a patrol route in a hallway.",
                "Mini-puzzle: filing system code to locate case numbers from E-05.",
                "Acquire E-06 and escape."
            ],
            ["Evidence E-06", "Unlock Building B-02 photo lab door logic"],
        ),
        MissionSpec(
            "M-03", "Final Mission: Saint Vesta Residence",
            "Deliver the twist via the cassette recording, then reveal Elena’s body.",
            "Ska climbs the oldest building in the city to find the room Elena prepared for him: a confession staged like a case file.",
            ["Reach the top room", "Play the cassette", "Find the hidden basement"],
            [
                "The building becomes quieter the higher Ska climbs.",
                "Dark Room: cassette plays; confession; gunshot.",
                "Basement: Elena’s body under a hanging light; final photo."
            ],
            ["Evidence E-09", "Ending headline + credits trigger"],
        ),
    ]

    # Buildings / Rooms
    b1 = BuildingSpec(
        "B-01", "Marrow Apartments",
        "A decaying mid-century block with water stains, old mailboxes, and a smell of wet plaster.",
        "Tone: claustrophobic, lived-in decay. Early-game breadcrumb trail.",
        rooms=[
            RoomSpec(
                "B1-R1", "Lobby & Mail Wall", "Ground Floor",
                "The player enters a narrow lobby with mailboxes, a broken intercom, and a faint perfume scent.",
                "Find the tenant name connected to Elena and acquire the first building clue.",
                layout=[
                    "Small rectangular lobby; left wall: mailboxes; back wall: stairwell door.",
                    "Right corner: broken vending machine; floor puddle reflecting neon from outside.",
                    "A barred window shows rain and the streetlamp flicker."
                ],
                interactables=[
                    "Mailbox panel (names scratched, some readable)",
                    "Intercom (broken but can be rewired)",
                    "Trash bin (contains torn envelope)",
                    "Vending machine slot (jammed, hides item behind)"
                ],
                puzzle=[
                    "Puzzle: Rewire Intercom",
                    "Player finds 3 loose wires (red/white/black). A note on the wall reads: “WHITE speaks, BLACK listens, RED lies.”",
                    "Correct wiring order (white -> speaker, black -> mic, red -> power) triggers a short audio glitch that reveals one tenant name before failing.",
                    "Reward: Tenant name: 'E. Vesta' and a partial apartment number."
                ],
                clue_rewards=[
                    "Evidence: E-02 Burned Matchbook (found behind the vending machine after fixing it)",
                    "Notebook update: “Saint Vesta Residence” appears as a recurring place-name."
                ],
                encounters=[
                    "Encounter: One 'Gray' thug looting mailboxes. He flees to stairwell if spotted.",
                    "Optional: chase micro-sequence; if failed, thug drops a key fragment."
                ],
                fail_states=[
                    "Wrong wiring triggers a fuse pop: lights out for 10 seconds (enemy can reposition).",
                    "If the thug escapes, player loses a small ammo pickup but still progresses."
                ],
                dialogue=[
                    "Ska: “Elena hated places like this… so why come here?”",
                    "Frank (radio call tutorial): “You’re walking into a graveyard of bad decisions, Ska.”"
                ],
                art_background=[
                    "Rain reflection shader on lobby floor puddles.",
                    "Mailboxes with scratched labels (handwritten look).",
                    "Flickering neon light spill through window bars.",
                    "Intercom panel: exposed wires, cracked plastic."
                ],
                art_characters=[
                    "Gray Thug variant A: hooded jacket, crowbar, nervous animations.",
                    "Ska idle: cigarette ember + coat flutter."
                ],
                vfx_sfx=[
                    "SFX: distant thunder, buzzing neon, intercom static.",
                    "VFX: brief spark burst on wrong wiring."
                ],
                exit_condition="Player learns apartment lead and opens the stairwell door.",
            ),
            RoomSpec(
                "B1-R2", "Janitor Closet", "Ground Floor",
                "A cramped closet packed with cleaning chemicals and a pinned maintenance map of the building.",
                "Acquire the master keyring that gates upper floors and future boiler access.",
                layout=[
                    "Tiny room: shelves left and right, mop bucket, corkboard with building map.",
                    "Small vent near ceiling with scratch marks.",
                ],
                interactables=[
                    "Corkboard building map (can be photographed)",
                    "Lockbox (requires code)",
                    "Vent cover (screwdriver needed)",
                    "Chemical shelf (can spill to create hazard)"
                ],
                puzzle=[
                    "Puzzle: Lockbox Code from Map",
                    "Corkboard map has 4 highlighted rooms circled with letters: A, R, C, H.",
                    "Player must count stair steps marked on the map between these rooms: A=3, R=1, C=4, H=2.",
                    "Enter 3142 to open the lockbox."
                ],
                clue_rewards=[
                    "Evidence: E-04 Maintenance Keyring",
                    "New capability: unlock Archive Office and Roof in B-01; unlock Boiler entry in B-02 later."
                ],
                encounters=[
                    "Encounter: None (puzzle/safe room).",
                    "Tension: occasional footsteps outside door (audio)."
                ],
                fail_states=[
                    "Three wrong codes triggers a loud alarm buzzer (spawns 1 thug in hallway outside upon exit)."
                ],
                dialogue=[
                    "Ska: “Maintenance maps… the city’s nervous system.”",
                    "Ska: “Someone planned this. Someone always plans this.”"
                ],
                art_background=[
                    "Grimy shelves with labeled bottles (bleach, solvent).",
                    "Corkboard with hand-drawn map and red string (foreshadow evidence board).",
                    "Vent scratches implying someone crawled through."
                ],
                art_characters=[
                    "No new characters required."
                ],
                vfx_sfx=[
                    "SFX: dripping pipe, distant elevator clunk.",
                ],
                exit_condition="Player picks up the keyring and returns to lobby/stairs.",
            ),
            RoomSpec(
                "B1-R3", "Apartment 1A (The Erased Face)", "Floor 1",
                "A small apartment with photos on the wall—every man’s face scratched out except one.",
                "Recover a photograph that proves Elena has been here, and identify the first 'husband' victim pattern.",
                layout=[
                    "Living area with couch, photo wall, kitchen nook.",
                    "Bedroom door half-open; a jewelry box on dresser.",
                    "Bathroom mirror cracked, reflecting neon lines."
                ],
                interactables=[
                    "Photo wall (inspect multiple photos)",
                    "Jewelry box (locked)",
                    "Kitchen drawer (contains razor blades)",
                    "Bathroom sink (blood stain can be revealed with UV)"
                ],
                puzzle=[
                    "Puzzle: Jewelry Box Ring Alignment",
                    "Jewelry box has 5 ring slots with symbols: Moon, Key, Eye, Rose, Knife.",
                    "Player must place rings found around the apartment into correct order using a clue poem under a photo:",
                    "“Moon watches. Key opens. Eye judges. Rose lies. Knife ends.”",
                    "Correct order opens hidden compartment."
                ],
                clue_rewards=[
                    "Evidence: E-03 Scratched Photograph",
                    "Clue: A torn diary page referencing 'Voss' and 'photo lab'."
                ],
                encounters=[
                    "Encounter: Two Grays enter from hallway after player grabs photo (trigger).",
                    "Combat beat: cramped room cover + throwable bottle."
                ],
                fail_states=[
                    "If player takes too long after trigger, enemies torch the photo wall (losing optional lore photos but not E-03)."
                ],
                dialogue=[
                    "Ska: “Who are you hiding from… Elena?”",
                    "Gray Thug: “Boss said bring the picture. Dead or alive.”"
                ],
                art_background=[
                    "Photo wall with scratched faces (distinct scratch patterns).",
                    "Kitchen clutter: cheap liquor, ashtray overflow.",
                    "Cracked mirror with reflected neon (simple parallax)."
                ],
                art_characters=[
                    "Gray Thug variant B: heavier build, uses baton.",
                    "Ska combat animations: short-range shove, pistol aim."
                ],
                vfx_sfx=[
                    "SFX: camera shutter click when evidence collected.",
                    "VFX: dust motes in sunlight strip."
                ],
                exit_condition="Defeat enemies and exit into hallway to Floor 2.",
            ),
            RoomSpec(
                "B1-R4", "Hallway Puzzle: The Broken Elevator", "Between Floors 1–2",
                "A vertical traversal challenge: elevator stuck; player must reroute power.",
                "Restore elevator power to access upper floors (or open an alternate stair route).",
                layout=[
                    "Narrow hallway with elevator door centered; breaker cabinet on right wall.",
                    "Stairwell door chained (optional to cut chain with tool found later)."
                ],
                interactables=[
                    "Breaker cabinet (mini-game)",
                    "Elevator panel (lights flicker)",
                    "Loose conduit (can be traced)",
                ],
                puzzle=[
                    "Puzzle: Circuit Load Balancing",
                    "Player flips 6 breakers; 3 must be ON to distribute load.",
                    "Clue is written on a maintenance sticker: “Never run HEAT, LIGHT, and LIFT at once.”",
                    "Solution: Turn OFF HEAT, keep LIGHT and LIFT, plus AUX.",
                    "Success powers elevator for a single ride."
                ],
                clue_rewards=[
                    "Progression gate unlocked: Floor 2 access.",
                    "Optional pickup: fuse (used later in B-02)."
                ],
                encounters=[
                    "Encounter: one ambush thug when cabinet opens (quick-time shove)."
                ],
                fail_states=[
                    "Wrong configuration triggers arc flash, minor damage, and resets puzzle."
                ],
                dialogue=[
                    "Ska: “Old buildings always tell the truth. They creak when you’re close.”"
                ],
                art_background=[
                    "Animated flicker light + subtle particle sparks around cabinet.",
                    "Elevator door grime + warning stickers."
                ],
                art_characters=[
                    "Gray Thug variant C: skinny, fast jab animation."
                ],
                vfx_sfx=[
                    "SFX: breaker clacks, elevator cable groan.",
                    "VFX: arc flash on failure."
                ],
                exit_condition="Elevator powers and carries player to Floor 2.",
            ),
            RoomSpec(
                "B1-R5", "Archive Office (The Numbers)", "Floor 2",
                "A cramped office filled with folders, typewriters, and a wall calendar full of crossed-out dates.",
                "Find the missing-husband case numbers and learn where police records are kept.",
                layout=[
                    "Desk with typewriter; file cabinets; wall calendar; safe under desk.",
                    "A torn city map on the wall with pins."
                ],
                interactables=[
                    "Typewriter (can type code to print)",
                    "File cabinets (locked without E-04 keyring)",
                    "Desk safe (requires 3-digit code)",
                    "City map pins (clue alignment)"
                ],
                puzzle=[
                    "Puzzle: Calendar Cipher",
                    "Calendar shows three dates circled: 10th, 14th, 22nd. Months: OCT, APR, FEB.",
                    "A sticky note reads: “Old rules: months are numbers, not names.”",
                    "Convert months to numbers: OCT=10, APR=4, FEB=2. Code becomes 1042.",
                    "Enter first 3 digits (104) into safe; fourth digit opens file cabinet latch (narrative excuse)."
                ],
                clue_rewards=[
                    "Evidence: E-05 Police Case Number List",
                    "Lead: “Records Room, Police HQ – sealed shelf 7.”"
                ],
                encounters=[
                    "Encounter: None inside office; threat is time pressure via sound of hallway search."
                ],
                fail_states=[
                    "If player triggers too much noise (3 loud actions), a thug bangs on door and forces a quick exit without optional loot."
                ],
                dialogue=[
                    "Ska: “Numbers don’t lie. People do.”",
                    "Ska: “Sealed shelf 7… still there. Unless they cleaned it.”"
                ],
                art_background=[
                    "Paper stacks, old folder labels, cigarette burns on desk.",
                    "Calendar with heavy pencil marks, dangling torn page corner.",
                ],
                art_characters=[
                    "No new characters required."
                ],
                vfx_sfx=[
                    "SFX: paper rustle, distant siren, typewriter key clicks.",
                ],
                exit_condition="Player exits to rooftop access door (unlocked via keyring).",
            ),
            RoomSpec(
                "B1-R6", "Rooftop Water Tank (Signal)", "Roof",
                "Open sky, rain, city skyline. A broken radio antenna and a water tank with graffiti.",
                "Use the rooftop vantage point to locate the next building and get a transient radio message from Elena.",
                layout=[
                    "Rooftop: water tank center; antenna on right; door back to stairwell.",
                    "Low wall edge (don’t fall; just visual)."
                ],
                interactables=[
                    "Antenna (repair)",
                    "Graffiti tag (contains clue)",
                    "Water tank hatch (locked)"
                ],
                puzzle=[
                    "Puzzle: Antenna Alignment",
                    "Player rotates antenna in 4 directions; the correct direction makes radio static reduce.",
                    "Clue: graffiti arrow points toward 'GULL ST' with a small compass mark.",
                    "Align antenna to Gull Street to get a short burst: a woman’s laugh + one word: “Voss.”"
                ],
                clue_rewards=[
                    "New lead unlocked: Building B-02 Gull Street Tenements.",
                    "Optional: water tank contains ammo if player found hatch key fragment earlier."
                ],
                encounters=[
                    "Encounter: One sniper-like enemy in adjacent building (telegraphed laser line). "
                    "Player must move between cover points to repair antenna."
                ],
                fail_states=[
                    "If hit during repair, progress resets by one step."
                ],
                dialogue=[
                    "Ska: “That voice… it can’t be.”",
                    "Ska: “Voss. Who the hell is Voss?”"
                ],
                art_background=[
                    "Parallax skyline layers with animated rain streaks.",
                    "Neon signs in distance; rooftop puddles.",
                    "Graffiti: 'GULL ST' readable."
                ],
                art_characters=[
                    "Enemy silhouette only (distant)."
                ],
                vfx_sfx=[
                    "SFX: wind, rain, antenna creak, radio static.",
                    "VFX: lightning flash occasional."
                ],
                exit_condition="Rooftop sequence completes and mission updates to travel to Gull Street.",
            ),
        ],
    )

    b2 = BuildingSpec(
        "B-02", "Gull Street Tenements",
        "Tight alley access, iron fire escapes, laundry lines. Feels like the city is watching you.",
        "Tone: paranoia, surveillance, hidden craftsmanship (photo lab). Mid-game escalation.",
        rooms=[
            RoomSpec(
                "B2-R1", "Courtyard & Fire Escape", "Exterior",
                "An enclosed courtyard with wet concrete and a fire escape ladder. Someone is photographing the entrance.",
                "Find and confront (or evade) the watcher; access the building without being photographed.",
                layout=[
                    "Courtyard center: drain and puddles; left: dumpster; right: fire escape ladder.",
                    "Upper balcony: shadow of a camera lens occasionally visible."
                ],
                interactables=[
                    "Dumpster (hides a torn letter)",
                    "Fire escape ladder (climb)",
                    "Courtyard light switch (can turn off to hide)",
                ],
                puzzle=[
                    "Puzzle: Light & Shadow Infiltration",
                    "Player must reach the door while staying in shadow when the camera flash occurs (every 6 seconds).",
                    "Optional solution: switch off courtyard light, which changes flash timing but makes enemies patrol closer."
                ],
                clue_rewards=[
                    "Clue: torn letter mentions “photo lab” and “wedding proofs.”",
                ],
                encounters=[
                    "Encounter: Watcher flees; player can chase for optional extra evidence node (no hard gate)."
                ],
                fail_states=[
                    "If photographed 3 times, two Grays spawn inside first hallway for immediate combat."
                ],
                dialogue=[
                    "Ska: “Someone’s documenting me. Same way I’m documenting her.”"
                ],
                art_background=[
                    "Wet laundry lines, dripping water, alley cats.",
                    "Camera flash light cone effect."
                ],
                art_characters=[
                    "Watcher NPC: slim silhouette with camera; non-combat."
                ],
                vfx_sfx=[
                    "SFX: camera shutter, distant train, dripping water.",
                    "VFX: flash bloom + brief afterimage."
                ],
                exit_condition="Player enters building through main door or fire escape window.",
            ),
            RoomSpec(
                "B2-R2", "Boiler Room (Heat)", "Basement",
                "A hot, loud boiler room with pipes like veins. A door is sealed by steam pressure.",
                "Use the fuse from B-01 and the maintenance keyring to stabilize pressure and open the sealed door.",
                layout=[
                    "Boiler center with valve wheels; pipe maze; control panel; sealed steel door back wall."
                ],
                interactables=[
                    "Pressure valves (3)",
                    "Control panel (missing fuse)",
                    "Thermometer gauge (clue)",
                ],
                puzzle=[
                    "Puzzle: Pressure Stabilization",
                    "Insert fuse (optional from B1-R4; if missing, player finds spare fuse behind loose panel).",
                    "Set valves to match gauge target: 70/40/30 (displayed as chalk marks on pipes).",
                    "When stable, door unlocks with a loud clang."
                ],
                clue_rewards=[
                    "Unlocks access to upper hallway route.",
                    "Optional clue: burn mark shaped like a ring on control panel."
                ],
                encounters=[
                    "Encounter: one heavy Gray with shield (uses steam bursts as cover)."
                ],
                fail_states=[
                    "If valves set wrong, steam burst reduces visibility and hurts player over time until reset."
                ],
                dialogue=[
                    "Ska: “Heat makes men honest. Or it makes them panic.”"
                ],
                art_background=[
                    "Pipe maze with animated steam vents.",
                    "Wet floor reflections, warning signs."
                ],
                art_characters=[
                    "Heavy Gray (shield/pipe wrench)."
                ],
                vfx_sfx=[
                    "SFX: boiler rumble, steam hiss, metal creaks.",
                    "VFX: thick steam particles."
                ],
                exit_condition="Sealed door opens; player progresses to Hallway of Rooms.",
            ),
            RoomSpec(
                "B2-R3", "Apartment 3C (Mara Voss’s Cache)", "Floor 3",
                "A careful, almost sterile apartment with labeled boxes. Someone lives like a librarian of crimes.",
                "Locate Mara Voss’s cache and obtain the lead to the hidden photo lab.",
                layout=[
                    "Main room with shelves of labeled evidence boxes.",
                    "Work desk with magnifying lens and glue.",
                    "Back room with locked wardrobe."
                ],
                interactables=[
                    "Evidence boxes (inspect labels)",
                    "Desk drawers (contain chemicals)",
                    "Wardrobe lock (pattern lock)"
                ],
                puzzle=[
                    "Puzzle: Label Sorting",
                    "Player finds 6 evidence box labels, 3 are counterfeit (ink slightly different).",
                    "Using magnifying lens, identify counterfeit labels (missing watermark).",
                    "Correct boxes reveal the photo lab key card hidden under false bottom."
                ],
                clue_rewards=[
                    "Key Item: Photo Lab Key Card",
                    "Notebook update: “Mara Voss collects wedding proofs.”"
                ],
                encounters=[
                    "Encounter: none until exit; then a hallway ambush.",
                ],
                fail_states=[
                    "If player opens wrong boxes 3 times, a silent alarm triggers and adds extra enemies in next encounter."
                ],
                dialogue=[
                    "Ska: “Voss isn’t a person. It’s a profession.”"
                ],
                art_background=[
                    "Clean shelves, labeled boxes, soft desk lamp cone.",
                    "Subtle unsettling detail: one box labeled “Ska, Adrian (future)”."
                ],
                art_characters=[
                    "No new characters in-room."
                ],
                vfx_sfx=[
                    "SFX: quiet room tone, clock tick, paper slide.",
                ],
                exit_condition="Player obtains key card and reaches the hidden lab entrance.",
            ),
            RoomSpec(
                "B2-R4", "Hidden Photo Lab (Proofs)", "Secret Room",
                "A darkroom with chemical trays, hanging photos, and a small tape fragment labeled for the 'final husband.'",
                "Recover the microcassette fragment and confirm Elena’s multiple identities through photo evidence.",
                layout=[
                    "Darkroom: red safety light; chemical trays; clothesline with photos.",
                    "Back wall: small safe; side door: exit to fire escape."
                ],
                interactables=[
                    "Photo line (inspect photos)",
                    "Chemical tray (can splash to blind enemy briefly)",
                    "Safe (combination)",
                ],
                puzzle=[
                    "Puzzle: Safe Combination from Photos",
                    "Three photos show wedding cakes with numbers on toppers: 7, 1, 9.",
                    "Player must order them by photo date stamps on the back (oldest to newest).",
                    "Combination becomes 197."
                ],
                clue_rewards=[
                    "Evidence: E-07 Audio Tape Fragment",
                    "Major reveal: Elena’s face appears with different surnames across decades."
                ],
                encounters=[
                    "Encounter: Two Grays burst in; low visibility fight under red light."
                ],
                fail_states=[
                    "If player knocks over all trays, room fills with smoke-like chemical fog (harder fight, but still winnable)."
                ],
                dialogue=[
                    "Ska: “Wedding photos… trophies.”",
                    "Ska: “She didn’t disappear. She moved on.”"
                ],
                art_background=[
                    "Red safety light with strong contrast shadows.",
                    "Hanging wet photos (subtle animation drip).",
                    "Chemical trays with reflection highlights."
                ],
                art_characters=[
                    "Gray Thug with gas mask variant (fits chemical room)."
                ],
                vfx_sfx=[
                    "SFX: dripping, paper flutter, muffled footsteps, chemical splash.",
                    "VFX: fog overlay if trays destroyed."
                ],
                exit_condition="Player secures tape fragment and escapes to prepare for Police HQ visit / final building.",
            ),
        ],
    )

    b3 = BuildingSpec(
        "B-03", "Saint Vesta Residence",
        "An old, almost sacred building: carved wood, antique lamps, silence that feels intentional.",
        "Tone: ritual, confession, finality. Late-game dread.",
        rooms=[
            RoomSpec(
                "B3-R1", "Entrance Hall (The Quiet)", "Ground Floor",
                "A grand hall with dust-covered furniture. It feels staged, like someone expected Ska.",
                "Enter the building and realize the space is prepared as a 'final case'.",
                layout=[
                    "Wide hall; staircase center; side doors left/right; chandelier above.",
                    "A guestbook on a pedestal."
                ],
                interactables=[
                    "Guestbook (names and dates)",
                    "Chandelier chain (can lower to access balcony later)",
                    "Side door locks (one opens, one sealed)"
                ],
                puzzle=[
                    "Puzzle: Guestbook Page Logic",
                    "Guestbook has pages stuck together. Player must warm them using a nearby lamp (heat) to separate without tearing.",
                    "Inside, a name appears repeatedly: Elena under multiple surnames.",
                    "This is a narrative puzzle + reinforcement of 'Many Names'."
                ],
                clue_rewards=[
                    "New narrative node: “This place is a ritual site.”"
                ],
                encounters=[
                    "Encounter: none; this is a calm-before-storm room."
                ],
                fail_states=[
                    "If player rips pages, they lose optional lore but not progression."
                ],
                dialogue=[
                    "Ska: “No footsteps… no life… like a church after the sermon.”"
                ],
                art_background=[
                    "Ornate wood panels, dusty cloth over furniture.",
                    "Floating dust particles in light beams."
                ],
                art_characters=[
                    "No new characters."
                ],
                vfx_sfx=[
                    "SFX: building creaks, distant wind through cracks.",
                ],
                exit_condition="Player unlocks stair access to upper floors.",
            ),
            RoomSpec(
                "B3-R2", "Caretaker’s Room (Ledger)", "Floor 1",
                "A small caretaker room with a ledger and a hidden latch diagram for the basement.",
                "Find the caretaker ledger and learn about Elena’s payments and the basement mechanism.",
                layout=[
                    "Bed, small desk, ledger, wall hooks with keys (some missing).",
                    "Carpet with a slightly raised corner."
                ],
                interactables=[
                    "Ledger (collect)",
                    "Carpet corner (reveal hatch latch sketch)",
                    "Key hooks (inspect missing outline)"
                ],
                puzzle=[
                    "Puzzle: Hook Shadow Pattern",
                    "Three key hooks cast shadows forming shapes at specific times (clock in room).",
                    "Player sets the clock hands to match a note: “Meet me at 3:15.”",
                    "Correct time reveals the latch combination drawn in shadow on the wall."
                ],
                clue_rewards=[
                    "Evidence: E-08 Caretaker Ledger",
                    "Basement latch pattern learned (used after cassette)."
                ],
                encounters=[
                    "Encounter: one Gray searches the room; stealth takedown tutorial option."
                ],
                fail_states=[
                    "If detected, Gray calls backup in hallway (adds 1 enemy later)."
                ],
                dialogue=[
                    "Ska: “Payments… she bought silence the way others buy flowers.”"
                ],
                art_background=[
                    "Old desk lamp, ledger with stained pages.",
                    "Carpet texture with raised corner hint."
                ],
                art_characters=[
                    "Gray stealth variant (flashlight)."
                ],
                vfx_sfx=[
                    "SFX: clock tick, pencil scratch when ledger opened.",
                ],
                exit_condition="Player secures ledger and continues toward top floors.",
            ),
            RoomSpec(
                "B3-R3", "Chapel Room (False Peace)", "Floor 2",
                "A room arranged like a private chapel—candles, photos, and wedding rings in a bowl.",
                "Survive an encounter and collect the last gating key to the top floor.",
                layout=[
                    "Altar-like table, rows of chairs, tall windows with rain streaks.",
                    "Side cabinet with locked drawer."
                ],
                interactables=[
                    "Ring bowl (inspect; count)",
                    "Candle stand (can be used as weapon)",
                    "Cabinet drawer (locked)"
                ],
                puzzle=[
                    "Puzzle: Ring Count Key",
                    "Cabinet lock requires a 2-digit number.",
                    "Clue: a note reads “All my husbands in one bowl.” Player counts rings (12).",
                    "Code is 12 to unlock the drawer containing the top-floor key."
                ],
                clue_rewards=[
                    "Key Item: Top Floor Key",
                    "Optional: ring with Ska’s initials (psychological hit)."
                ],
                encounters=[
                    "Encounter: mini-boss 'Gray Captain' with stun baton; uses chair rows as cover.",
                ],
                fail_states=[
                    "If player is stunned 3 times, they drop ammo and must recover it mid-fight."
                ],
                dialogue=[
                    "Gray Captain: “She said you’d come. She always knows.”",
                    "Ska: “Then tell her I’m done being predictable.”"
                ],
                art_background=[
                    "Candle light flicker; ring bowl shine highlights.",
                    "Chapel chairs, rain streak window parallax."
                ],
                art_characters=[
                    "Gray Captain: distinct coat, baton, confident stance."
                ],
                vfx_sfx=[
                    "SFX: candle crackle, rain on glass, baton electric hum.",
                    "VFX: small sparks on baton hits."
                ],
                exit_condition="Player defeats mini-boss and unlocks access to top floor.",
            ),
            RoomSpec(
                "B3-R4", "Top Hallway (The Last Steps)", "Top Floor",
                "A long hallway with doors that are all locked except the final room. The silence is heavy.",
                "Use all collected evidence to open the final room and commit to the truth.",
                layout=[
                    "Long corridor, peeling wallpaper, one functioning lamp at far end.",
                    "Final door with multiple locks."
                ],
                interactables=[
                    "Multi-lock door (requires evidence nodes)",
                    "Wall portraits (faces covered)"
                ],
                puzzle=[
                    "Puzzle: Evidence Board Gate",
                    "The final door unlocks only if the player has connected the core nodes:",
                    "E-01 -> B-01, E-05 -> Police HQ, E-06 -> Many Names, E-07 -> Final Husband.",
                    "Mechanically: a UI check; narratively: Ska accepts the conclusion."
                ],
                clue_rewards=[
                    "Access to the Dark Room (cassette)."
                ],
                encounters=[
                    "Encounter: none; this is a pacing breath."
                ],
                fail_states=[
                    "If missing a key evidence, objective instructs player to backtrack."
                ],
                dialogue=[
                    "Ska: “Every step up feels like a step into a confession booth.”"
                ],
                art_background=[
                    "Hallway lamp cone, deep shadows, dust drifting.",
                    "Portraits with cloth over faces."
                ],
                art_characters=[
                    "No new characters."
                ],
                vfx_sfx=[
                    "SFX: footsteps echo, lamp buzz, distant building groan.",
                ],
                exit_condition="Door unlocks; player enters the Dark Room.",
            ),
            RoomSpec(
                "B3-R5", "Dark Room (Cassette Confession)", "Top Floor – Final Room",
                "Pitch-black room. A desk in the center. A cassette recorder and a note: “To my final husband.”",
                "Play the cassette, receive the twist, trigger the ending sequence.",
                layout=[
                    "Nearly empty black room; desk center; chair; cassette deck.",
                    "One weak light source that turns on only after PLAY is pressed."
                ],
                interactables=[
                    "Cassette deck (PLAY)",
                    "Note on desk (read)",
                ],
                puzzle=[
                    "Interaction set-piece (no fail puzzle): player must choose to press PLAY.",
                    "Optional: allow player to walk away; but the objective remains until they commit."
                ],
                clue_rewards=[
                    "Evidence: E-09 Cassette Confession (automatic after PLAY)."
                ],
                encounters=[
                    "Encounter: none (pure narrative)."
                ],
                fail_states=[
                    "None. The emotional 'fail state' is the truth."
                ],
                dialogue=[
                    "Cassette (Elena): “Hello, Ska.”",
                    "Cassette (Elena): “If you’re hearing this, then you’ve solved my little mystery.”",
                    "Cassette (Elena): “My name isn’t really Elena. I’ve worn many names. Many rings. Many husbands.”",
                    "Cassette (Elena): “I killed them, Ska. Every single one. No conspiracy. Just me.”",
                    "Cassette (Elena): “I am sorry… because you were the only one I ever loved.”",
                    "Cassette: (revolver being loaded) … (gunshot)."
                ],
                art_background=[
                    "Ultra-dark palette; desk silhouette; subtle dust in beam.",
                    "Cassette deck detailed sprite with moving tape wheels."
                ],
                art_characters=[
                    "No new characters on-screen; focus on Ska silhouette."
                ],
                vfx_sfx=[
                    "SFX: cassette click, tape hiss, voice filter, revolver cylinder, gunshot reverb.",
                    "VFX: slight screen shake on gunshot; fade to near-black."
                ],
                exit_condition="After cassette ends, a distant noise triggers; basement sequence objective appears.",
            ),
            RoomSpec(
                "B3-R6", "Hidden Basement (Final Photograph)", "Basement",
                "A hidden room under a single hanging light: Elena’s body, revolver in hand, a final photograph.",
                "Confirm the ending, deliver final headline/credits.",
                layout=[
                    "Small basement room, single light, concrete walls, drain.",
                    "Body placement near center; photo near hand."
                ],
                interactables=[
                    "Photograph (inspect)",
                    "Basement door (exit triggers credits)",
                ],
                puzzle=[
                    "No puzzle. This is a controlled ending tableau.",
                    "Optional: player can replay cassette in inventory, but it changes nothing."
                ],
                clue_rewards=[
                    "Ending trigger: headline and credits."
                ],
                encounters=[
                    "Encounter: none."
                ],
                fail_states=[
                    "None."
                ],
                dialogue=[
                    "Ska (silent): (no dialogue — the only 'line' is silence).",
                    "UI Headline: “THE CITY'S GREATEST DETECTIVE SOLVES HIS LAST CASE.”"
                ],
                art_background=[
                    "Hanging light cone with dust particles.",
                    "Concrete stains, subtle water drip animation.",
                    "Photograph close-up asset (Ska + Elena in happier time)."
                ],
                art_characters=[
                    "Elena body sprite (tastefully framed; noir shadow to avoid gore)."
                ],
                vfx_sfx=[
                    "SFX: faint drip, distant building settling, low rumble.",
                    "VFX: slow camera zoom on photograph; fade to black."
                ],
                exit_condition="Player inspects photograph; screen fades; credits roll.",
            ),
        ],
    )

    police_hq_rooms = [
        RoomSpec(
            "HQ-R1", "Police HQ – Lobby & Metal Detector", "Police HQ",
            "The old workplace. Cold fluorescent lighting. Posters about procedure.",
            "Enter without alerting Internal Affairs and reach the records wing.",
            layout=[
                "Reception desk, metal detector, hallway to records wing.",
                "Security camera cone (simple stealth)."
            ],
            interactables=[
                "Reception logbook (fake sign-in)",
                "Camera switchboard (temporarily loops feed)",
                "Metal detector tray (stash gun briefly)"
            ],
            puzzle=[
                "Puzzle: Sign-In Disguise",
                "Player must choose correct alias from three options based on a name tag found earlier in B-01 (optional).",
                "Correct choice avoids suspicion; wrong choice triggers a short chase but still possible to escape."
            ],
            clue_rewards=[
                "Access to Records Room."
            ],
            encounters=[
                "Encounter: stealth-lite; no combat inside HQ unless player fails badly."
            ],
            fail_states=[
                "If caught by IA twice, player is forced out and must re-enter via service door (adds extra puzzle)."
            ],
            dialogue=[
                "Ska: “This place taught me rules. Elena taught me exceptions.”"
            ],
            art_background=[
                "Fluorescent flicker, reflective tile floor, police signage.",
            ],
            art_characters=[
                "Internal Affairs Officer: clean suit, badge, slow patrol route.",
                "Receptionist: bored animation loop."
            ],
            vfx_sfx=[
                "SFX: fluorescent hum, distant radio chatter, footsteps.",
            ],
            exit_condition="Player reaches records corridor.",
        ),
        RoomSpec(
            "HQ-R2", "Police HQ – Records Room (Sealed Shelf 7)", "Police HQ",
            "Rows of file cabinets and a sealed shelf section. The air smells like paper and dust.",
            "Use E-05 case numbers to retrieve Elena’s marriage record copy.",
            layout=[
                "Large room with shelves; sealed shelf area; desk lamp; copy machine."
            ],
            interactables=[
                "Shelf 7 lock (numeric)",
                "Copy machine (prints evidence copy)",
                "Case ledger (search UI)"
            ],
            puzzle=[
                "Puzzle: Filing System",
                "Player enters 3 case numbers from E-05 into the ledger to receive shelf slots.",
                "Then uses slot initials to form lock code: e.g., S-K-A (19-11-1) => 19111.",
                "Simplify in implementation: choose a fixed code referenced in E-05 notes."
            ],
            clue_rewards=[
                "Evidence: E-06 Wedding Certificate Copy",
                "Narrative beat: Elena’s signature is different, but the handwriting matches."
            ],
            encounters=[
                "Encounter: time pressure—IA patrol passes doorway every 20 seconds."
            ],
            fail_states=[
                "If caught, player must drop copies and flee; can retry later but loses optional extra file."
            ],
            dialogue=[
                "Ska: “Paper remembers what people forget.”"
            ],
            art_background=[
                "Tall shelves, file labels, dust motes in lamp cone.",
            ],
            art_characters=[
                "IA Officer patrol silhouette in doorway.",
            ],
            vfx_sfx=[
                "SFX: paper slide, stamp thud, copy machine whirr.",
            ],
            exit_condition="Player secures E-06 and escapes HQ.",
        ),
    ]

    # Return all content
    return {
        "story_overview": story_overview,
        "characters": characters,
        "pillars": pillars,
        "evidence": evidence,
        "missions": missions,
        "buildings": [b1, b2, b3],
        "police_hq_rooms": police_hq_rooms,
    }


# -----------------------------
# Rendering (DOCX writing)
# -----------------------------

def render_gdd(doc: Document, content: Dict[str, object]) -> None:
    doc.add_heading("Detective Ska — Game Design Document", level=1)
    doc.add_paragraph("Genre: 2D noir detective / mystery with light combat and puzzle rooms.")
    doc.add_paragraph("Engine: Unity (2D).")
    doc.add_paragraph("Primary gameplay: investigation, environmental puzzles, evidence linking, room-based progression.")

    doc.add_heading("1. Story Overview", level=2)
    doc.add_paragraph(str(content["story_overview"]))

    doc.add_heading("2. Tone & Themes", level=2)
    add_bullets(doc, [
        "Noir grief and obsession: the case is the only thing keeping Ska upright.",
        "Truth as punishment: every clue hurts, but ignorance is worse.",
        "Love as the perfect disguise: Elena’s affection is the mask that fooled everyone (including Ska)."
    ])

    doc.add_heading("3. Main Characters", level=2)
    for name, desc in content["characters"]:
        doc.add_heading(name, level=3)
        doc.add_paragraph(desc)

    doc.add_heading("4. Core Gameplay Pillars", level=2)
    add_numbered(doc, list(content["pillars"]))

    doc.add_heading("5. Evidence System", level=2)
    doc.add_paragraph(
        "Evidence items are collected in rooms and connected on an Evidence Board. "
        "Progress gates are implemented as required evidence-node connections."
    )
    doc.add_heading("5.1 Evidence Catalog", level=3)
    ev_rows = [("Code", "Name / Found In / What it does")]
    for e in content["evidence"]:
        ev_rows.append((e.code, f"{e.name}\nFound in: {e.found_in}\n{e.description}\nUnlocks: {e.unlocks}"))
    add_kv_table(doc, ev_rows)

    doc.add_heading("6. Missions (High-Level)", level=2)
    for m in content["missions"]:
        doc.add_heading(f"{m.code} — {m.name}", level=3)
        doc.add_paragraph(m.overview)
        add_kv_table(doc, [
            ("Purpose", m.purpose),
        ])
        doc.add_paragraph("Player Goals:")
        add_bullets(doc, m.player_goals)
        doc.add_paragraph("Key Beats:")
        add_numbered(doc, m.key_beats)
        doc.add_paragraph("Deliverables (what design must ship):")
        add_bullets(doc, m.deliverables)
        if m.notes:
            doc.add_paragraph("Notes:")
            add_bullets(doc, m.notes)

    add_section_break(doc)

    doc.add_heading("7. Level Design — Room-by-Room Specification", level=2)
    doc.add_paragraph(
        "Structure: Three main investigation buildings, each composed of rooms that deliver "
        "one concrete clue plus one gameplay beat. Each room below includes objective, puzzle, encounters, "
        "dialogue snippets, and art/sound tasks."
    )

    # Building rooms
    for b in content["buildings"]:
        doc.add_heading(f"{b.code} — {b.name}", level=2)
        doc.add_paragraph(b.address_style)
        doc.add_paragraph(b.tone)

        for r in b.rooms:
            render_room(doc, r)
            doc.add_paragraph("")  # spacing

        add_section_break(doc)

    # Police HQ section (intermission)
    doc.add_heading("8. Intermission Area — Police HQ (Room Specs)", level=2)
    doc.add_paragraph(
        "Police HQ is an intermission mission space used to convert building clues into official records. "
        "It plays as stealth-lite: avoid Internal Affairs, solve filing puzzles, retrieve marriage documents."
    )
    for r in content["police_hq_rooms"]:
        render_room(doc, r)
        doc.add_paragraph("")

    add_section_break(doc)

    doc.add_heading("9. Cutscenes & Key Dialogue (Selected)", level=2)
    doc.add_heading("9.1 Bar Scene (Black Raven) — Core Dialogue", level=3)
    add_dialogue_block(doc, [
        "Frank Ramirez: “You look like a man trying to drown a ghost.”",
        "Ska: “I’m trying to forget a name.”",
        "Frank Ramirez: “Then don’t ask me to say it out loud.”",
        "Stranger (overheard): “She’s been seen near Marrow again. Same perfume. Same smile.”",
        "Ska (quiet): “Elena…”",
        "(Optional: Bar fight sequence triggers if the player threatens the Stranger or refuses to leave.)",
    ])

    doc.add_heading("9.2 Cassette Confession — Required Lines", level=3)
    doc.add_paragraph(
        "This is the twist delivery. Keep the room dark. Make the player press PLAY themselves."
    )
    add_dialogue_block(doc, [
        "Cassette: “Hello, Ska.”",
        "Cassette: “If you’re hearing this, then you’ve solved my little mystery.”",
        "Cassette: “My name isn’t really Elena… I’ve worn many names.”",
        "Cassette: “And many husbands.”",
        "Cassette: “I killed them, Ska. Every single one.”",
        "Cassette: “I am sorry… because you were the only one I ever loved.”",
        "Cassette: (revolver loaded) … (gunshot).",
    ])

    doc.add_heading("10. Art Bible (Practical Tasks)", level=2)
    doc.add_heading("10.1 Background Artist Checklist", level=3)
    add_bullets(doc, [
        "Define 3 building palettes: B-01 mildew green, B-02 sickly sodium yellow, B-03 antique brown/black.",
        "Create reusable prop set: mailboxes, intercom, file cabinets, desk lamps, photo frames, cassette deck, pipes.",
        "Create rain/puddle tiles and parallax skyline layers for rooftop scenes.",
        "Make “evidence close-up” assets: receipts, photos, diary pages, certificates, ledger, cassette label."
    ])
    doc.add_heading("10.2 Character Artist Checklist", level=3)
    add_bullets(doc, [
        "Ska: idle, walk, examine, pistol aim, shove, hurt, climb ladder, interact animations.",
        "Frank Ramirez: bartender idle, wiping glass, leaning forward, gesture talk.",
        "Gray thugs: 3 variants + captain (baton). Include flashlight stealth variant.",
        "Watcher NPC: camera pose, run/flee, idle."
    ])
    doc.add_heading("10.3 VFX/SFX Checklist", level=3)
    add_bullets(doc, [
        "Neon flicker loop; rain ambience; distant sirens.",
        "Intercom static & electrical sparks; breaker arc flash.",
        "Boiler steam particles; chemical fog variant.",
        "Cassette audio filter + gunshot reverb + subtle screen shake."
    ])

    doc.add_heading("11. Ending", level=2)
    doc.add_paragraph(
        "After the cassette ends, Ska follows a distant noise to a hidden basement room. "
        "Under a single hanging light lies Elena’s body, revolver in hand, a final photograph nearby. "
        "Ska kneels beside her. For the first time in the game, he says nothing. Fade to black.\n\n"
        "Final headline:\n"
        "“THE CITY'S GREATEST DETECTIVE SOLVES HIS LAST CASE.”\n"
        "Credits roll."
    )


def render_room(doc: Document, r: RoomSpec) -> None:
    doc.add_heading(f"{r.code} — {r.name} ({r.floor})", level=3)
    doc.add_paragraph(r.summary)

    add_kv_table(doc, [
        ("Objective", r.objective),
        ("Exit Condition", r.exit_condition),
    ])

    doc.add_paragraph("Layout Overview:")
    add_bullets(doc, r.layout)

    doc.add_paragraph("Key Interactables:")
    add_bullets(doc, r.interactables)

    doc.add_paragraph("Puzzle / Progression:")
    add_numbered(doc, r.puzzle)

    doc.add_paragraph("Clues & Rewards:")
    add_bullets(doc, r.clue_rewards)

    doc.add_paragraph("Encounters:")
    add_bullets(doc, r.encounters)

    if r.fail_states:
        doc.add_paragraph("Fail States / Consequences:")
        add_bullets(doc, r.fail_states)

    if r.dialogue:
        doc.add_paragraph("Dialogue Snippets:")
        add_dialogue_block(doc, r.dialogue)

    doc.add_paragraph("Art Tasks — Background:")
    add_bullets(doc, r.art_background)

    doc.add_paragraph("Art Tasks — Characters:")
    add_bullets(doc, r.art_characters)

    doc.add_paragraph("VFX / SFX Notes:")
    add_bullets(doc, r.vfx_sfx)


# -----------------------------
# Entry point
# -----------------------------

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Generate Detective Ska GDD docx.")
    p.add_argument(
        "--out",
        default="Detective_Ska_GDD.docx",
        help="Output DOCX file path (default: Detective_Ska_GDD.docx)",
    )
    p.add_argument(
        "--no-toc",
        action="store_true",
        help="Do not insert a Table of Contents field.",
    )
    return p.parse_args()


def main() -> None:
    args = parse_args()
    content = build_content()

    doc = Document()
    add_title_page(
        doc,
        "Detective Ska",
        "Game Design Document (Room-by-room, puzzle + art tasks + dialogue)",
    )
    if not args.no_toc:
        add_toc(doc)
    render_gdd(doc, content)
    doc.save(args.out)
    print(f"Saved: {args.out}")
    print("Tip: Open the DOCX in Word and update fields to populate the Table of Contents.")


if __name__ == "__main__":
    main()

